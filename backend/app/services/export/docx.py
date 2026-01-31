import re
import uuid
import asyncio
from pathlib import Path
from typing import Dict, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.core.logger import get_logger
from app.core.config import settings

class DOCXExporter:
    """Génère des fichiers DOCX professionnels avec images intégrées"""
    
    def __init__(self):
        self.logger = get_logger("export.docx")
        self.output_dir = Path(settings.DOCS_DIR)
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def _clean_text(self, text: str) -> str:
        """
        Supprime les caractères de contrôle XML invalides qui font 
        échouer la sauvegarde du document DOCX.
        """
        if not text:
            return ""
        # Garde uniquement les caractères valides pour le XML de Word
        return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)

    async def generate_docx(self, note_data: Dict, filename: Optional[str] = None) -> tuple[str, int]:
        try:
            if filename is None:
                # Utilisation de .hex pour un nom de fichier plus propre
                filename = f"export_{uuid.uuid4().hex}.docx"
            
            if not filename.endswith('.docx'):
                filename += '.docx'
            
            file_path = self.output_dir / filename
            doc = Document()

            # --- Style par défaut ---
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)

            # 1) Titre Principal
            title = self._clean_text(note_data.get('title', 'Notes de cours'))
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            content = note_data.get('content', '')
            lines = content.split('\n')
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue

                # --- GESTION DES IMAGES ---
                image_match = re.match(r'!\[.*?\]\((.*?)\)', line)
                if image_match:
                    img_path = Path(image_match.group(1))
                    if img_path.exists():
                        try:
                            # Insertion de l'image centrée
                            doc.add_picture(str(img_path), width=Inches(5.0))
                            last_paragraph = doc.paragraphs[-1]
                            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            # Légende optionnelle
                            # Dans la gestion des images, remplace la partie caption par :
                            try:
                                caption_text = "Illustration technique"
                                caption = doc.add_paragraph(caption_text)
                                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                # On tente d'appliquer le style, sinon on formate manuellement
                                try:
                                    caption.style = 'Caption'
                                except KeyError:
                                    run = caption.runs[0] if caption.runs else caption.add_run(caption_text)
                                    run.font.size = Pt(9)
                                    run.italic = True
                            except Exception as e:
                                self.logger.warning(f"Légende non ajoutée : {e}")
                        except Exception as e:
                            self.logger.error(f"Erreur insertion image DOCX: {e}")
                    continue

                # --- NETTOYAGE DE LA LIGNE ---
                line = self._clean_text(line)

                # --- GESTION DES BLOCS D'ALERTE (> ⚠️) ---
                if line.startswith('> '):
                    p = doc.add_paragraph(line.replace('> ', ''), style='Quote')
                    if p.runs:
                        run = p.runs[0]
                        run.font.color.rgb = RGBColor(200, 0, 0)
                        run.bold = True
                    continue

                # --- GESTION DU MARKDOWN CLASSIQUE ---
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith('1. ') or re.match(r'^\d+\.', line):
                    text = re.sub(r'^\d+\.\s', '', line)
                    doc.add_paragraph(text, style='List Number')
                else:
                    doc.add_paragraph(line)

            # Sauvegarde asynchrone pour ne pas bloquer l'event loop
            await asyncio.to_thread(doc.save, str(file_path))
            
            file_size = file_path.stat().st_size
            self.logger.info(f"✅ DOCX généré avec succès: {file_path}")
            return str(file_path), file_size

        except Exception as e:
            self.logger.error(f"❌ Erreur critique lors de la génération DOCX: {e}")
            raise e

# Instance et alias pour l'export
docx_exporter = DOCXExporter()
generate_docx = docx_exporter.generate_docx